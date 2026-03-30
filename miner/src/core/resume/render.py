from __future__ import annotations

from abc import ABC, abstractmethod
from pathlib import Path
import tempfile
from typing import TYPE_CHECKING, Union
from pdflatex import PDFLaTeX
from src.infrastructure.log.logging import get_logger

logger = get_logger(__name__)

if TYPE_CHECKING:
    from src.core.resume.resume import Resume, ResumeItem


class ResumeRender(ABC):
    @abstractmethod
    def render(self, resume: Resume) -> Union[str, bytes]: ...


class TextResumeRenderer(ResumeRender):
    def render(self, resume: Resume) -> str:
        to_return = ""
        to_return += f"Email: {resume.email}\n\n" if resume.email else ""
        to_return += f"Core skills: {', '.join(resume.skills)}\n\n" if resume.skills else ""

        # Check if resume has education entries before rendering
        if resume.education:
            to_return += "Education:\n"
            for ed in resume.education:
                to_return += f"   - {ed}\n"
            to_return += "\n"

        # Check if resume has awards entries before rendering
        if resume.awards:
            to_return += "Awards:\n"
            for aw in resume.awards:
                to_return += f"   - {aw}\n"
            to_return += "\n"

        for item in resume.items:
            to_return += f"{item.title} : {item.start_date.strftime('%B, %Y')} - {item.end_date.strftime('%B, %Y')}\n"

            if len(item.frameworks) != 0:
                to_return += f"Frameworks: {', '.join(skill.skill_name for skill in item.frameworks)}\n"

            for bullet in item.bullet_points:
                to_return += f"   - {bullet}\n"
            if item is not resume.items[-1]:
                to_return += "\n"
            to_return += "\n"

        return to_return


LATEX_SPECIAL_CHARS = {
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
    "\\": r"\textbackslash{}",
}


def latex_escape(text: str) -> str:
    return "".join(LATEX_SPECIAL_CHARS.get(c, c) for c in text)


class PDFRenderer(ResumeRender):
    """
    This class extends the resumeRender and allows for the user to
    convert the given LaTeX resume to PDF format for unfamiliar users
    """

    def render(self, resume) -> bytes:
        pdf = b""

        with tempfile.TemporaryDirectory() as tmpdir:
            tex_path = Path(tmpdir) / "helper.tex"
            tex_path.write_text(
                resume.export(ResumeLatexRenderer()),
                encoding="utf-8"
            )

            pdfLaTeX = PDFLaTeX.from_texfile(str(tex_path))
            pdfLaTeX.set_interaction_mode()
            try:
                pdf, _, _ = pdfLaTeX.create_pdf(keep_pdf_file=True)
            except FileNotFoundError:
                logger.error("There is a problem with your TeX Live installation.\n"
                             "Run which pdflatex in your terminal to see if TeX Live is properly installed.\n"
                             "If nothing is shown, visit https://www.tug.org/texlive/ for installation instructions")

        return pdf

# --- Renderer ---


class ResumeLatexRenderer(ResumeRender):
    """
    Generates a stable LaTeX resume using the Jake Gutierrez template.
    """

    prefix = r"""
\documentclass[letterpaper,11pt]{article}

\usepackage{latexsym}
\usepackage[empty]{fullpage}
\usepackage{titlesec}
\usepackage[usenames,dvipsnames]{color}
\usepackage{verbatim}
\usepackage{enumitem}
\usepackage[hidelinks]{hyperref}
\usepackage{fancyhdr}
\usepackage[english]{babel}
\usepackage{tabularx}
\input{glyphtounicode}

% Formatting
\pagestyle{fancy}
\fancyhf{}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}

% Clean margins
\usepackage[margin=0.7in]{geometry}

\raggedbottom
\raggedright
\setlength{\tabcolsep}{0in}

% Section formatting
\titleformat{\section}{
  \vspace{-4pt}\scshape\raggedright\large
}{}{0em}{}[\color{black}\titlerule \vspace{-5pt}]

\pdfgentounicode=1

% --- Custom Commands ---
\newcommand{\resumeItem}[1]{
  \item\small{
    {#1 \vspace{-2pt}}
  }
}

\newcommand{\resumeSubheading}[4]{
  \vspace{-2pt}\item
    \begin{tabular*}{0.97\textwidth}[t]{l@{\extracolsep{\fill}}r}
      \textbf{#1} & #2 \\
      \textit{\small#3} & \textit{\small #4} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeProjectHeading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \small#1 & #2 \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeItemListStart}{\begin{itemize}}
\newcommand{\resumeItemListEnd}{\end{itemize}\vspace{-5pt}}
\newcommand{\resumeSubHeadingListStart}{\begin{itemize}[leftmargin=0.15in, label={}]}
\newcommand{\resumeSubHeadingListEnd}{\end{itemize}}
"""

    def render(self, resume: Resume) -> str:
        tex = [self.prefix, r"\begin{document}"]

        # --- Header ---
        tex.append(r"\begin{center}")

        name = latex_escape(resume.name) if resume.name else "Your Name"
        tex.append(rf"\textbf{{\Huge \scshape {name}}} \\ \vspace{{1pt}}")

        contact_parts = []
        if resume.location:
            contact_parts.append(rf"\small {latex_escape(resume.location)}")
        if resume.email:
            email = latex_escape(resume.email)
            contact_parts.append(
                rf"\href{{mailto:{email}}}{{\underline{{{email}}}}}")
        if resume.linkedin:
            linkedin = resume.linkedin
            href = linkedin if linkedin.startswith("http") else f"https://{linkedin}"
            contact_parts.append(
                rf"\href{{{latex_escape(href)}}}{{\underline{{LinkedIn}}}}")
        if resume.github:
            github = latex_escape(resume.github)
            contact_parts.append(
                rf"\href{{https://github.com/{github}}}{{\underline{{GitHub}}}}")

        if contact_parts:
            tex.append(" $|$ ".join(contact_parts) + r" \\")

        tex.append(r"\end{center}")
        tex.append("")

        # --- Education ---
        if resume.education:
            tex.append(r"\section{Education}")
            tex.append(r"\resumeSubHeadingListStart")
            for ed in resume.education:
                if isinstance(ed, dict):
                    title = latex_escape(ed.get("title", ""))
                    start = ed.get("start") or ""
                    end = ed.get("end") or ""
                    date_str = f"{start} -- {end}" if start and end else (start or end)
                    if date_str:
                        tex.append(rf"\resumeProjectHeading{{\textbf{{{title}}}}}{{{latex_escape(date_str)}}}")
                    else:
                        tex.append(rf"  \item \small{{{title}}}")
                else:
                    tex.append(rf"  \item \small{{{latex_escape(str(ed))}}}")
            tex.append(r"\resumeSubHeadingListEnd")
            tex.append("")

        # --- Awards ---
        if resume.awards:
            tex.append(r"\section{Awards}")
            tex.append(r"\resumeSubHeadingListStart")
            for aw in resume.awards:
                if isinstance(aw, dict):
                    title = latex_escape(aw.get("title", ""))
                    start = aw.get("start") or ""
                    end = aw.get("end") or ""
                    date_str = f"{start} -- {end}" if start and end else (start or end)
                    if date_str:
                        tex.append(rf"\resumeProjectHeading{{\textbf{{{title}}}}}{{{latex_escape(date_str)}}}")
                    else:
                        tex.append(rf"  \item \small{{{title}}}")
                else:
                    tex.append(rf"  \item \small{{{latex_escape(str(aw))}}}")
            tex.append(r"\resumeSubHeadingListEnd")
            tex.append("")

        # --- Projects ---
        if resume.items:
            tex.append(r"\section{Projects}")
            tex.append(r"\resumeSubHeadingListStart")

            for item in resume.items:
                title = latex_escape(item.title)
                start = item.start_date.strftime('%B %Y') if item.start_date else ""
                end = item.end_date.strftime('%B %Y') if item.end_date else ""
                date = f"{start} -- {end}" if start or end else ""

                if item.frameworks:
                    frameworks = ", ".join(
                        latex_escape(f.skill_name) for f in item.frameworks)
                    title_tex = rf"\textbf{{{title}}} $|$ \emph{{{frameworks}}}"
                else:
                    title_tex = rf"\textbf{{{title}}}"

                tex.append(rf"\resumeProjectHeading{{{title_tex}}}{{{date}}}")
                tex.append(r"\resumeItemListStart")
                for b in item.bullet_points:
                    tex.append(rf"\resumeItem{{{latex_escape(b)}}}")
                tex.append(r"\resumeItemListEnd")

            tex.append(r"\resumeSubHeadingListEnd")
            tex.append("")

        # --- Technical Skills ---
        if resume.skills:
            skills = ", ".join(latex_escape(s) for s in resume.skills)
            tex.append(r"\section{Technical Skills}")
            tex.append(r"\begin{itemize}[leftmargin=0.15in, label={}]")
            tex.append(rf"  \item \small{{{skills}}}")
            tex.append(r"\end{itemize}")

        tex.append(r"\end{document}")

        return "\n".join(tex)


class DocxResumeRenderer(ResumeRender):
    """
    Generates a .docx resume using python-docx.
    Mirrors the section structure of the LaTeX renderer:
    Header → Education → Awards → Projects → Technical Skills
    """

    def render(self, resume: Resume) -> bytes:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io

        doc = Document()

        # ── Narrow margins ──────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # ── Default style ────────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        def _add_section_heading(text: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            # Bottom border to mimic \titlerule
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "000000")
            pBdr.append(bottom)
            pPr.append(pBdr)
            return p

        def _add_dated_row(title: str, date_str: str):
            """One paragraph: bold title on left, italic date on right via tab stop."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            # Right tab stop at page width minus margins
            from docx.oxml import OxmlElement as _OX
            from docx.oxml.ns import qn as _qn
            from docx.shared import Twips
            pPr = p._p.get_or_add_pPr()
            tabs = _OX("w:tabs")
            tab = _OX("w:tab")
            tab.set(_qn("w:val"), "right")
            tab.set(_qn("w:pos"), str(int(Twips(9360))))  # ~6.5 in
            tabs.append(tab)
            pPr.append(tabs)
            run_title = p.add_run(title)
            run_title.bold = True
            p.add_run("\t")
            run_date = p.add_run(date_str)
            run_date.italic = True
            run_date.font.size = Pt(10)
            return p

        def _add_bullet(text: str):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(text)
            run.font.size = Pt(10)
            return p

        # ── Header ───────────────────────────────────────────────────────────
        name = resume.name or "Your Name"
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_after = Pt(2)
        run_name = p_name.add_run(name)
        run_name.bold = True
        run_name.font.size = Pt(18)

        contact_parts = []
        if resume.location:
            contact_parts.append(resume.location)
        if resume.email:
            contact_parts.append(resume.email)
        if resume.linkedin:
            contact_parts.append(resume.linkedin)
        if resume.github:
            contact_parts.append(f"github.com/{resume.github}")

        if contact_parts:
            p_contact = doc.add_paragraph(" | ".join(contact_parts))
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_contact.paragraph_format.space_before = Pt(0)
            p_contact.paragraph_format.space_after = Pt(4)
            for run in p_contact.runs:
                run.font.size = Pt(10)

        # ── Education ────────────────────────────────────────────────────────
        if resume.education:
            _add_section_heading("Education")
            for ed in resume.education:
                if isinstance(ed, dict):
                    title = ed.get("title", "")
                    start = ed.get("start") or ""
                    end = ed.get("end") or ""
                    date_str = f"{start} \u2013 {end}" if start and end else (start or end)
                    if date_str:
                        _add_dated_row(title, date_str)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(title)
                else:
                    p = doc.add_paragraph()
                    p.add_run(str(ed))

        # ── Awards ───────────────────────────────────────────────────────────
        if resume.awards:
            _add_section_heading("Awards")
            for aw in resume.awards:
                if isinstance(aw, dict):
                    title = aw.get("title", "")
                    start = aw.get("start") or ""
                    end = aw.get("end") or ""
                    date_str = f"{start} \u2013 {end}" if start and end else (start or end)
                    if date_str:
                        _add_dated_row(title, date_str)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(title)
                else:
                    p = doc.add_paragraph()
                    p.add_run(str(aw))

        # ── Projects ─────────────────────────────────────────────────────────
        if resume.items:
            _add_section_heading("Projects")
            for item in resume.items:
                start = item.start_date.strftime("%B %Y") if item.start_date else ""
                end = item.end_date.strftime("%B %Y") if item.end_date else ""
                date_str = f"{start} \u2013 {end}" if start and end else (start or end)

                if item.frameworks:
                    fw_str = ", ".join(f.skill_name for f in item.frameworks)
                    title_str = f"{item.title} | {fw_str}"
                else:
                    title_str = item.title

                _add_dated_row(title_str, date_str)
                for bullet in item.bullet_points:
                    _add_bullet(bullet)

        # ── Technical Skills ─────────────────────────────────────────────────
        if resume.skills:
            _add_section_heading("Technical Skills")
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.add_run(", ".join(resume.skills)).font.size = Pt(10)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
